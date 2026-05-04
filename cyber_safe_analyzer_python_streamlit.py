# CyberSafe Class Analyzer — Python Streamlit version
# Run:
# 1) pip install streamlit python-docx
# 2) streamlit run app.py

import re
import io
from collections import Counter
from html import escape

import streamlit as st
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="CyberSafe Class Analyzer", page_icon="🛡️", layout="wide")

RISK_LEXICON = {
    "high": [
        "убью", "өлтірем", "побью", "изобью", "зарежу", "сдохни", "умри", "угрожаю",
        "на три буквы", "посылает", "иди на", "шантаж", "суицид", "повешусь"
    ],
    "medium": [
        "тупая", "тупой", "дура", "дурак", "идиот", "лох", "заткнись", "не спрашивала",
        "не у тебя", "зачем ты", "ты сказала", "не брала", "как она у тебя", "гонишь",
        "врала", "обман", "неге алдың", "алдың"
    ],
    "low": [
        "ахах", "хаха", "😂", "🤣", "ммм", "ого", "копеец", "емма", "ема", "🤦", "лол"
    ],
    "self_defense": [
        "я не", "мен емес", "не брала", "я знаю не", "не тупая", "мен алмадым",
        "не виновата", "я не могла", "я не делала"
    ],
    "system": [
        "сообщения и звонки защищены", "добавил(-а)", "создал(-а) группу", "изменил(-а)",
        "закрепил(-а)", "теперь вы админ", "исчезающие сообщения", "удалил", "удалила",
        "присоединился", "настройки группы", "вы удалили", "данное сообщение удалено",
        "изображение отсутствует", "аудиофайл отсутствует", "видео отсутствует",
        "видеозаметка опущена", "<сообщение изменено>"
    ]
}

CHAT_PATTERNS = [
    re.compile(r"^\[(\d{1,2}\.\d{1,2}\.\d{4}),\s*(\d{1,2}:\d{2}(?::\d{2})?)\]\s*([^:]+):\s*(.*)$"),
    re.compile(r"^(\d{1,2}\.\d{1,2}\.\d{4}),\s*(\d{1,2}:\d{2}(?::\d{2})?)\s*-\s*([^:]+):\s*(.*)$"),
    re.compile(r"^(\d{1,2}/\d{1,2}/\d{2,4}),\s*(\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\s*-\s*([^:]+):\s*(.*)$"),
    re.compile(r"^\[(\d{1,2}/\d{1,2}/\d{2,4}),\s*(\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM|am|pm)?)\]\s*([^:]+):\s*(.*)$"),
]


def normalize_name(name: str) -> str:
    name = name.replace("~", "").replace("🍫", "")
    name = re.sub(r"5\s?б", "", name, flags=re.IGNORECASE)
    name = re.sub(r"\s+", " ", name)
    return name.strip()


def parse_chat(text: str):
    messages = []
    unmatched = []
    lines = text.splitlines()

    for line in lines:
        clean = line.replace("\u200e", "").replace("\u200f", "").strip()
        if not clean:
            continue

        match = None
        for pattern in CHAT_PATTERNS:
            match = pattern.match(clean)
            if match:
                break

        if match:
            date, time, sender, message = match.groups()
            messages.append({
                "date": date.replace("/", "."),
                "time": time.strip(),
                "sender": normalize_name(sender),
                "message": (message or "").strip(),
                "raw": clean,
            })
        elif messages:
            messages[-1]["message"] += "\n" + clean
        else:
            unmatched.append(clean)

    return messages, unmatched


def contains_any(text: str, words: list[str]):
    lower = text.lower()
    return [word for word in words if word.lower() in lower]


def is_system_or_media(message: str) -> bool:
    lower = message.lower()
    return any(word.lower() in lower for word in RISK_LEXICON["system"])


def classify_message(message: str):
    high = contains_any(message, RISK_LEXICON["high"])
    medium = contains_any(message, RISK_LEXICON["medium"])
    low = contains_any(message, RISK_LEXICON["low"])
    self_defense = contains_any(message, RISK_LEXICON["self_defense"])

    if high:
        return "High", "Қорлау / қауіп / вербалды агрессия", "Жеке сөйлесу, ата-анамен байланыс, қайталанса әкімшілік шара", high
    if medium:
        return "Medium", "Агрессивті коммуникация / қысым", "Чатта тоқтату, жеке түсіндіру, конфликтті офлайн шешу", medium
    if self_defense:
        return "Medium", "Өзін қорғау реакциясы", "Оқушымен жеке сөйлесу, эмоционалдық қолдау көрсету", self_defense
    if low:
        return "Low", "Күлкі / мазақ болуы мүмкін", "Контекстті бақылау, мазаққа айналса ескерту", low
    return "None", "Қалыпты", "Бақылау қажет емес", []


def filter_keywords(messages, keywords):
    result = []
    for msg in messages:
        lower = msg["message"].lower()
        if any(k.lower() in lower for k in keywords):
            result.append(msg)
    return result


def unique_senders(items):
    return ", ".join(sorted(set(item["sender"] for item in items if item.get("sender"))))


def quote_items(items, limit=5):
    quotes = []
    for item in items[:limit]:
        text = re.sub(r"\s+", " ", item["message"])[:160]
        quotes.append(f"“{text}”")
    return quotes


def identify_situations(messages):
    situations = []

    correction = filter_keywords(messages, ["замаз", "не брала", "зачем ты", "как она у тебя", "ты сказала"])
    if len(correction) >= 2:
        situations.append({
            "title": "Замазка оқиғасы",
            "participants": unique_senders(correction),
            "summary": "Жоғалған затқа байланысты бір оқушыға бірнеше сұрақ қойылып, айыптау сипаты байқалған.",
            "evaluation": "Толық буллинг емес, бірақ микро-буллинг және топтық қысым белгісі бар.",
            "quotes": quote_items(correction),
        })

    admin = filter_keywords(messages, ["админ", "староста", "зам", "слежу", "поряд"])
    if len(admin) >= 3:
        situations.append({
            "title": "Админ/староста рөліне байланысты жағдай",
            "participants": unique_senders(admin),
            "summary": "Оқушылар арасында чаттағы басқару рөлі, админ болу және жауапкершілікке байланысты пікірталас болған.",
            "evaluation": "Буллинг емес, бірақ лидерлік пен билікке таласу арқылы конфликт тудыруы мүмкін.",
            "quotes": quote_items(admin),
        })

    aggression = filter_keywords(messages, ["посылает", "на три буквы", "иди на", "тупая", "тупой", "дура", "дурак"])
    if len(aggression) >= 1:
        situations.append({
            "title": "Вербалды агрессия белгілері",
            "participants": unique_senders(aggression),
            "summary": "Чатта немесе чат арқылы сыныптағы дөрекі сөйлеу, боқтау туралы хабарламалар кездеседі.",
            "evaluation": "Бұл ең маңызды қауіп факторы. Қайталанса буллингке айналуы мүмкін.",
            "quotes": quote_items(aggression),
        })

    return situations


def analyze_chat(raw_text: str):
    parsed, unmatched = parse_chat(raw_text)
    text_messages = [m for m in parsed if m["message"] and not is_system_or_media(m["message"])]

    enriched = []
    for msg in text_messages:
        risk, category, recommendation, keywords = classify_message(msg["message"])
        item = dict(msg)
        item.update({
            "risk": risk,
            "category": category,
            "recommendation": recommendation,
            "keywords": keywords,
        })
        enriched.append(item)

    participants = sorted(set(m["sender"] for m in text_messages if m["sender"]))
    risky = [m for m in enriched if m["risk"] != "None"]
    high = [m for m in risky if m["risk"] == "High"]
    medium = [m for m in risky if m["risk"] == "Medium"]
    low = [m for m in risky if m["risk"] == "Low"]

    sender_counts = Counter(m["sender"] for m in enriched)
    active_threshold = max(3, int(len(text_messages) * 0.03))
    active = [(name, count) for name, count in sender_counts.most_common() if count >= active_threshold]

    score = len(high) * 3 + len(medium) * 2 + len(low)
    risk_level = "🟢 Төмен"
    if len(high) > 0 or len(medium) >= 5 or score >= 12:
        risk_level = "🟡 Орта"
    if len(high) >= 3 or len(medium) >= 15 or score >= 40:
        risk_level = "🔴 Жоғары"

    return {
        "parsed": parsed,
        "unmatched": unmatched,
        "text_messages": text_messages,
        "enriched": enriched,
        "participants": participants,
        "risky": risky,
        "high": high,
        "medium": medium,
        "low": low,
        "active": active,
        "situations": identify_situations(enriched),
        "risk_level": risk_level,
    }


def quote_line(items):
    if not items:
        return "- Нақты мәтіндік хабарлама анықталмады."
    lines = []
    for item in items[:6]:
        text = re.sub(r"\s+", " ", item["message"])[:160]
        lines.append(f"- “{text}” ({item['sender']}, {item['date']} {item['time']})")
    return "\n".join(lines)


def generate_report(a):
    situation_text = "Нақты конфликт ситуациялары анықталмады."
    if a["situations"]:
        blocks = []
        for i, s in enumerate(a["situations"], 1):
            quotes = "\n".join(f"- {q}" for q in s["quotes"])
            blocks.append(
                f"{i}) {s['title']}\n"
                f"Қатысушылар: {s['participants'] or 'нақты анықталмады'}.\n"
                f"Не болды: {s['summary']}\n"
                f"Бағалау: {s['evaluation']}\n"
                f"Мысалдар:\n{quotes}"
            )
        situation_text = "\n\n".join(blocks)

    active_text = ", ".join(f"{name} ({count})" for name, count in a["active"]) or "анықталмады"

    return f"""5 «Б» СЫНЫБЫНЫҢ WHATSAPP ЧАТЫ БОЙЫНША ПСИХОЛОГИЯЛЫҚ ЕСЕП

1. Жалпы мәлімет

Бұл есеп WhatsApp чатындағы мәтіндік хабарламаларды талдау негізінде жасалды. Жүйелік хабарламалар, аудио, видео және сурет орынбасарлары есепке алынбады. Негізгі назар оқушылардың өзара қарым-қатынасына, сөйлеу мәдениетіне және кибербуллинг белгілеріне аударылды.

Қатысушылар саны: {len(a['participants'])}. Мәтіндік хабарламалар саны: {len(a['text_messages'])}. Белсенді қатысушылар саны: {len(a['active'])}.

2. Жалпы сипаттама

Чаттың негізгі мазмұны оқу процесіне және ұйымдастыруға байланысты. Оқушылар үй тапсырмасын сұрайды, бір-біріне жауап береді, ақпарат алмасады. Сонымен қатар чат эмоционалды және белсенді сипатта.

3. Анықталған мәселелер

Қауіпті хабарламалар саны: {len(a['risky'])}. Жоғары қауіп: {len(a['high'])}, орта қауіп: {len(a['medium'])}, төмен қауіп: {len(a['low'])}.

Жоғары қауіп деңгейіндегі хабарламалар:
{quote_line(a['high'])}

Орта қауіп деңгейіндегі хабарламалар:
{quote_line(a['medium'])}

Төмен қауіп деңгейіндегі хабарламалар:
{quote_line(a['low'])}

4. Нақты ситуациялар

{situation_text}

5. Психологиялық анализ

Чаттың жалпы атмосферасы аралас сипатта: негізгі бөлігі қалыпты, бірақ кей жерлерде конфликт, қысым және агрессивті сөйлеу элементтері байқалады. Сынып белсенді, тез жауап беретін және эмоционалды ұжым ретінде көрінеді.

Белсенді қатысушылар: {active_text}. Бұл оқушылар чат атмосферасына көбірек әсер етеді.

6. Қауіп деңгейі

Жалпы қауіп деңгейі: {a['risk_level']}. Бұл деңгей кейбір хабарламаларда дөрекі сөйлеу, топтық қысым, мазақ немесе өзін қорғау реакциялары байқалуына байланысты қойылды.

7. Ұсыныстар

Бірінші, чат ережесін бекіту қажет: қорлау, айыптау, мазақ ету және конфликтті көпшілік чатта талқылауға болмайды.

Екінші, қауіпті немесе дөрекі хабарлама жазған оқушылармен жеке сөйлесу керек.

Үшінші, қысымға түскен немесе өзін қорғауға мәжбүр болған оқушымен жеке сөйлесіп, эмоционалдық жағдайын анықтау қажет.

Профилактика ретінде «Онлайн қарым-қатынас мәдениеті» тақырыбында сынып сағатын өткізу ұсынылады.

8. Қорытынды

Жалпы алғанда, чат оқу және ұйымдастыру мақсатында қолданылып отыр. Дегенмен кейбір хабарламаларда дөрекі сөйлеу, қысым және конфликт белгілері бар. Қазіргі жағдайда толық жүйелі буллинг анықталмады, бірақ ерте кезеңдегі қауіп белгілері бар.
"""


def create_word(report_text: str):
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = document.add_heading("CyberSafe Class Analyzer", level=1)
    title.alignment = 1

    for block in report_text.split("\n"):
        if block.strip():
            document.add_paragraph(block)
        else:
            document.add_paragraph("")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------- UI ----------------
st.title("🛡️ CyberSafe Class Analyzer")
st.caption("WhatsApp TXT чатын жүктеңіз немесе мәтінді вставить етіңіз. Жүйе қазақша психологиялық есеп жасап, Word файлын береді.")

left, right = st.columns([1, 1])

with left:
    st.subheader("1) Чатты енгізу")
    uploaded_file = st.file_uploader("WhatsApp .txt файл жүктеу", type=["txt"])
    pasted_text = st.text_area("Немесе TXT мәтінді осында вставить етіңіз", height=300, placeholder="[12.11.2025, 22:36:55] Аты: Хабарлама...")

    raw_text = ""
    if uploaded_file is not None:
        raw_text = uploaded_file.read().decode("utf-8", errors="ignore")
        st.success(f"Файл жүктелді: {uploaded_file.name}")
    elif pasted_text.strip():
        raw_text = pasted_text
        st.success("Мәтін қабылданды")

    analyze_button = st.button("🔍 Анализ жасау", use_container_width=True, type="primary")

with right:
    st.subheader("2) Нәтиже")

    if analyze_button:
        if not raw_text.strip():
            st.error("Алдымен TXT файл жүктеңіз немесе мәтінді вставить етіңіз.")
        else:
            analysis = analyze_chat(raw_text)
            report = generate_report(analysis)
            word_file = create_word(report)

            st.session_state["analysis"] = analysis
            st.session_state["report"] = report
            st.session_state["word_file"] = word_file

    if "analysis" in st.session_state:
        analysis = st.session_state["analysis"]
        report = st.session_state["report"]

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Қатысушылар", len(analysis["participants"]))
        col2.metric("Мәтіндік хабарлама", len(analysis["text_messages"]))
        col3.metric("Қауіпті хабарлама", len(analysis["risky"]))
        col4.metric("Қауіп деңгейі", analysis["risk_level"])

        st.download_button(
            "📄 Word отчет жүктеу",
            data=st.session_state["word_file"],
            file_name="cybersafe_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.download_button(
            "📝 TXT отчет жүктеу",
            data=report.encode("utf-8"),
            file_name="cybersafe_report.txt",
            mime="text/plain",
            use_container_width=True,
        )

        st.subheader("Дайын есеп")
        st.text_area("Есеп мәтіні", report, height=520)

        st.subheader("Қауіпті хабарламалар кестесі")
        risky_table = [
            {
                "Дата": m["date"],
                "Уақыт": m["time"],
                "Қатысушы": m["sender"],
                "Хабарлама": m["message"][:180],
                "Белгі": m["category"],
                "Қауіп": m["risk"],
                "Ұсыныс": m["recommendation"],
            }
            for m in analysis["risky"][:100]
        ]
        st.dataframe(risky_table, use_container_width=True)
    else:
        st.info("Анализ нәтижесі осы жерде шығады.")
